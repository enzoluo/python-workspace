# -*- coding: UTF-8 -*-
# author:enzoluo
import datetime
import re

import requests
from bs4 import BeautifulSoup
from docx import Document

# 获取访问地址

# 城市故事
url = "http://www.jianshu.com/c/2ba824b6ed53?utm_medium=index-collections&utm_source=desktop"
headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:61.0) Gecko/20100101 Firefox/61.0"}

document_title = "城市故事"
# 谈谈情说说爱
# url = 'http://www.jianshu.com/c/GQ5FAs'
# document_title = "谈谈情说说爱"
# 互联网 IT
# url = 'https://www.jianshu.com/c/V2CqjW'
# document_title = "互联网 IT"
res = requests.get(url, headers=headers)
res.encoding = "utf-8"
soup = BeautifulSoup(res.text, 'lxml')

# print(soup)

# 声明接下来要是用的变量，变量的名字要声明清楚

# 接下来大概要获取这些数据：文章的名字，访问的连接，文章的描述，还有文章的正文内容

titles = []
links = []
descriptions = []
contents = []

title_soup = soup.select("#list-container > ul > li > div > a")
descritpion_soup = soup.select("#list-container > ul > li > div > p")

# 正则表达式
re_title = re.compile(r'target="_blank">(.*?)</a>')
re_description = re.compile(r'<p class="abstract">[\s\D]* (.*?)\n')
re_link = re.compile(r'href="(.*?)"')


def selectData(re_temp, html, dataList_param):
    for x in html:
        temp_d = re.findall(re_temp, str(x))
        temp_d = str(temp_d).strip("['").strip("']").strip("<br/>").strip("""\xa0""")
        dataList_param.append(temp_d)


def selectLink(re_temp, html):
    for x in html:
        link_temp = re.findall(re_temp, str(x))
        link_temp = "http://www.jianshu.com" + str(link_temp).strip("['").strip("']")
        links.append(link_temp)


# 筛选数据
selectData(re_title, title_soup, titles)
selectData(re_description, descritpion_soup, descriptions)
selectLink(re_link, title_soup)


# 对每一个文章具体内容的获取

def selectContent(re_temp, links_param, contents_param):
    for link in links_param:
        article_url = link
        article_res = requests.get(article_url,headers=headers)
        article_res.encoding = "utf-8"
        article_soup = BeautifulSoup(article_res.text, 'lxml')
        content_soup = article_soup.select("body > div.note > div.post > div.article > div.show-content p")
        temp_content = []
        selectData(re_temp, content_soup, temp_content)
        temp_str = ""
        for x in temp_content:
            temp_str = temp_str + str(x) + '\n'
        contents_param.append(temp_str)


reContent = re.compile(r'<p>(.*?)</p>')
selectContent(reContent, links, contents)

# 把所有的数据写入到一个文件中
articles = []
for title, description, content in zip(titles, descriptions, contents):
    temp = {
        "title": "标题：" + title,
        "description": description,
        "content": content
    }
    articles.append(temp)

document = Document()
document.styles['Normal'].font.name = u'宋体'
document.add_heading(document_title, 0)
for article in articles:
    document.add_heading(article.get('title'), 1)
    document.add_heading(article.get('description'), 2)
    document.add_paragraph(article.get('content'))
# 获取当前时间
now = datetime.datetime.now().strftime('%Y_%m_%d_%H_%M_%S')

document.save('C:/Users/10947/Desktop/美文' + now + '.docx')
