# -*- coding: UTF-8 -*-
# author:enzoluo
import docx
from docx import*
from docx.shared import *
from docx.oxml.ns import qn
import os
import datetime


# create a new document
document = Document()
document.add_heading("古诗",0)
document.add_heading("春江花月夜",1)


p1 = document.add_paragraph('床前明月光，')
p2 = document.add_paragraph('疑似地上霜。')
p3 = document.add_paragraph('举头望明月，')
p4 = document.add_paragraph('低头思故乡。')
document.save('C:/Users/user/Desktop/demo.docx')

# rs = os.path.exists('C:/Users/user/Desktop/demo.docx')
# print(rs)
now = datetime.datetime.now().strftime('%Y_%m_%d_%H_%M_%S')
print(now)

# # add some level title
# document.add_heading("Demo docx")
# document.add_heading(u"二级标题",1)
# document.add_heading(u"二级标题",2)
#
# #add text
# paragragh = document.add_paragraph(u'添加了文本')
#
# #set font size
# run = paragragh.add_run(u'设置字号')
# run.font.size=Pt(24)

# #set font family
# run = paragragh.add_run('Set Font,')
# run.font.name = 'Consolas'
#
# # set chiese font family
# run = paragragh.add_run('Set Font,')
# run.font.name=u'微软雅黑'
# r = run._element
# r.rPr.rFonts.set(qn('w:eastAsia'),u'微软雅黑')
#
# # set font I
# run = paragragh.add_run(u'斜体、')
#
# # set font height
# document.save('demo.docx')
#


